import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import re

# 创建一个新的Word文档
doc = docx.Document('软件学报2016年排版样例模版.docx')

# 清除模板中的内容，但保留样式
while len(doc.paragraphs) > 0:
    p = doc.paragraphs[0]._element
    p.getparent().remove(p)
    p._p = p._element = None

# 读取提取的PDF全文
with open('all_text.txt', 'r', encoding='utf-8') as file:
    content = file.read()

# 分割内容为标题和段落
sections = re.split(r'(^#{1,6}\s+.*$)', content, flags=re.MULTILINE)

# 处理标题和段落
current_level = 0
for i, section in enumerate(sections):
    if section.strip() == '':
        continue
    
    # 处理标题
    if section.startswith('#'):
        level = len(re.match(r'^(#+)', section).group(1))
        title = section.strip('#').strip()
        
        # 添加标题段落
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if level == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # 设置标题样式
        run = p.add_run(title)
        run.bold = True
        
        if level == 1:  # 主标题
            run.font.size = Pt(16)
        elif level == 2:  # 一级标题
            run.font.size = Pt(14)
        elif level == 3:  # 二级标题
            run.font.size = Pt(12)
        else:  # 三级及以下标题
            run.font.size = Pt(10.5)
        
        current_level = level
    
    # 处理正文段落
    else:
        paragraphs = section.strip().split('\n')
        for para in paragraphs:
            if para.strip() == '':
                continue
                
            # 处理表格
            if para.startswith('|') and '|' in para[1:]:
                # 这里简单处理表格，实际情况可能需要更复杂的逻辑
                rows = [row for row in para.split('\n') if row.strip()]
                if len(rows) >= 2:  # 至少有表头和分隔行
                    # 解析表头
                    header = rows[0].strip().split('|')
                    header = [cell.strip() for cell in header if cell.strip()]
                    
                    # 创建表格
                    table = doc.add_table(rows=1, cols=len(header))
                    # 移除不存在的样式名称，使用默认样式
                    
                    # 添加表头
                    header_cells = table.rows[0].cells
                    for i, text in enumerate(header):
                        header_cells[i].text = text
                    
                    # 添加数据行
                    for row in rows[2:]:  # 跳过表头和分隔行
                        cells = row.strip().split('|')
                        cells = [cell.strip() for cell in cells if cell.strip()]
                        if cells:
                            row_cells = table.add_row().cells
                            for i, text in enumerate(cells):
                                if i < len(row_cells):
                                    row_cells[i].text = text
                continue
                
            # 处理普通段落（含引用角标）
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            # 分割文本与引用角标（支持连续多个角标如[60][61]）
            parts = re.split(r'(\[\d+\](?:\[\d+\])*)', para)
            for part in parts:
                if not part:
                    continue
                # 检测是否为引用角标（支持多个连续角标）
                if re.fullmatch(r'\[\d+\](?:\[\d+\])*', part):
                    # 拆分多个连续角标并逐个处理
                    sub_parts = re.findall(r'\[\d+\]', part)
                    for sub_part in sub_parts:
                        run = p.add_run(sub_part.strip('[]'))
                        run.font.superscript = True  # 设置上标
                        run.font.size = Pt(10.5)  # 角标字号与正文一致
                        run.font.kerning = Pt(1)  # 调整字间距避免重叠
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10.5)  # 正文字体大小10.5磅
                    run.font.name = '宋体'  # 正文默认宋体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 显式设置东亚字体

# 添加符合《软件学报2016年排版样例模板》的参考文献列表
ref_heading = doc.add_paragraph('参考文献')
ref_heading.style = doc.styles['标题2']  # 使用模板标题样式
ref_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 创建带边框的表格（模板要求）
ref_table = doc.add_table(rows=1, cols=2)
# 设置表格边框（模板要求）
table_style = doc.styles['Table Grid']
ref_table.style = table_style
# 设置表头样式（模板要求加粗、10.5磅字）
header_cells = ref_table.rows[0].cells
header_run1 = header_cells[0].paragraphs[0].add_run('角标')
header_run1.font.bold = True
header_run1.font.size = Pt(10.5)
header_run2 = header_cells[1].paragraphs[0].add_run('文献信息')
header_run2.font.bold = True
header_run2.font.size = Pt(10.5)
# 示例文献（需根据实际翻译内容补充，可扩展此列表）
# 从论文PDF中提取的实际引用文献（示例，需用户补充完整）
ref_rows = [
    ('1', 'Achilles: Efficient TEE-Assisted BFT Consensus via Rollback Resilient Recovery'),
    ('60', '示例文献60：[作者]. [文献标题]. [期刊/会议], [年份]:[页码].')  # 按《软件学报》格式补充
]
for idx, info in ref_rows:
        row_cells = ref_table.add_row().cells
        # 数据行样式（严格模板要求：宋体、10.5磅字）
        # 角标列样式
        idx_run = row_cells[0].paragraphs[0].add_run(idx)
        idx_run.font.size = Pt(10.5)
        idx_run.font.name = '宋体'
        idx_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 文献信息列样式
        info_run = row_cells[1].paragraphs[0].add_run(info)
        info_run.font.size = Pt(10.5)
        info_run.font.name = '宋体'
        info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 显式设置东亚字体为宋体
# 保存文档
doc.save('Achilles.docx')

print("文档已生成：翻译.docx")