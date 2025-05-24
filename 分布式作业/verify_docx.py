import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# 打开生成的文档
try:
    doc = docx.Document('Achilles.docx')
    print('成功打开Achilles.docx')
except Exception as e:
    print(f'打开文档失败: {e}')
    exit(1)

# 验证引用角标（上标）
found_superscript = False
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.superscript:
            found_superscript = True
            break
    if found_superscript:
        break

# 验证参考文献列表
ref_found = False
ref_table_found = False
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '参考文献':
        ref_found = True
        # 检查后续是否有表格
        if i+1 < len(doc.paragraphs):
            next_element = doc._element.body[i+1]
            if next_element.tag.endswith('tbl'):
                # 遍历表格找到与next_element对应的表格对象
                for table in doc.tables:
                    if table._element == next_element:
                        ref_table = table
                header_cells = ref_table.rows[0].cells
                # 验证表头文本和样式
                # 验证表头文本和样式
                header_text_ok = header_cells[0].text.strip() == '角标' and header_cells[1].text.strip() == '文献信息'
                print(f'表头文本验证: {header_text_ok}')
                header_bold_ok = header_cells[0].paragraphs[0].runs[0].font.bold and header_cells[1].paragraphs[0].runs[0].font.bold
                print(f'表头加粗验证: {header_bold_ok}')
                header_font_size_ok = header_cells[0].paragraphs[0].runs[0].font.size == Pt(10.5) and header_cells[1].paragraphs[0].runs[0].font.size == Pt(10.5)
                print(f'表头字号验证: {header_font_size_ok} (实际字号: {header_cells[0].paragraphs[0].runs[0].font.size}, {header_cells[1].paragraphs[0].runs[0].font.size})')
                # 验证数据行样式（字体和字号）
                data_font_ok = False
                if len(ref_table.rows) > 1:
                    data_cell = ref_table.rows[1].cells[1]
                    data_run = data_cell.paragraphs[0].runs[0]
                    # 正确获取东亚字体设置（解决font.name可能不准确的问题）
                    eastasia_font = data_run._element.rPr.rFonts.get(qn('w:eastAsia'))
                    data_font_ok = (eastasia_font == '宋体' and data_run.font.size == Pt(10.5))
                    print(f'数据行字体验证: {data_font_ok} (实际字体: {data_run.font.name}, 实际字号: {data_run.font.size})')
                else:
                    print('数据行不存在，无法验证字体')
                # 综合所有检查项
                if header_text_ok and header_bold_ok and header_font_size_ok and data_font_ok:
                    ref_table_found = True
                    print('所有表格验证条件通过')
                else:
                    print('表格验证未通过，以上条件存在不满足项')
        break

# 输出验证结果
print('\n=== 验证结果 ===')
print('引用角标存在: ' + ('是' if found_superscript else '否'))
print('参考文献标题存在: ' + ('是' if ref_found else '否'))
print('参考文献表格格式正确: ' + ('是' if ref_table_found else '否'))